# renders content.py into a Word .docx laid out to the IEEE Computer Society
# proceedings look (US Letter, Times 10pt, numbered references).
# needs python-docx: pip install python-docx
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import content

HERE = Path(__file__).resolve().parent


def set_cols(section, num=2, space_twips=450):
    """turn a section into num columns. the IEEE CS format uses two with a
    5/16 inch (0.3125 in = 450 twips) gap between them."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))


def page_setup(section):
    # IEEE Computer Society 8.5x11 proceedings: print area 6.875 x 8.875 in,
    # which works out to these margins.
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin, section.bottom_margin = Inches(1.0), Inches(1.125)
    section.left_margin, section.right_margin = Inches(0.8125), Inches(0.8125)

# sections IEEE leaves without a roman numeral
UNNUMBERED = {"Acknowledgment", "Data and Code Availability"}

# {key} citations become [n] in order of the reference list
NUM = {k: i + 1 for i, (k, _) in enumerate(content.REFERENCES)}


def cite_sub(text):
    def repl(m):
        keys = re.findall(r"\{([a-z]+)\}", m.group(0))
        return "[" + ", ".join(str(NUM[k]) for k in keys) + "]"
    return re.sub(r"(\{[a-z]+\})+", repl, text)


def add_body(doc, text, size=10, justify=True):
    p = doc.add_paragraph()
    run = p.add_run(cite_sub(text))
    run.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.2)  # IEEE indents each para
    return p


def add_figure(doc):
    f = content.FIGURE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(HERE / f["file"]), width=Inches(3.0))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run("Fig. 1.  ")
    run.bold = True
    run.font.size = Pt(8)
    cap.add_run(f["caption"]).font.size = Pt(8)


def add_ranking_table(doc):
    t = content.TABLE
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("TABLE I")
    r.bold = True
    r.font.size = Pt(8)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run(t["caption"])
    rr.italic = True
    rr.font.size = Pt(8)
    tbl = doc.add_table(rows=1, cols=len(t["columns"]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, c in enumerate(t["columns"]):
        run = tbl.rows[0].cells[i].paragraphs[0].add_run(c)
        run.bold = True
        run.font.size = Pt(8)
    for row in t["rows"]:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val)).font.size = Pt(8)
    # force a fixed layout with an explicit grid, otherwise the narrow rank
    # column autofits down to nothing when Word or LibreOffice renders it
    _fix_widths(tbl, [720, 2016, 864])  # twips: 0.5in, 1.4in, 0.6in


def _fix_widths(tbl, widths_twips):
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for col, w in zip(grid.findall(qn("w:gridCol")), widths_twips):
        col.set(qn("w:w"), str(w))
    for row in tbl.rows:
        for cell, w in zip(row.cells, widths_twips):
            cell.width = Inches(w / 1440)


def render():
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Times New Roman"
    normal.size = Pt(10)

    # title and author span the full width (single column)
    page_setup(doc.sections[0])

    # title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(content.TITLE)
    r.bold = True
    r.font.size = Pt(24)

    # author with the required high-school affiliation
    a = content.AUTHOR
    au = doc.add_paragraph()
    au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    au.add_run(a["name"] + "\n" + "\n".join(a["affiliation"]))

    # everything from here down is two columns, IEEE style
    body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    page_setup(body_sec)
    set_cols(body_sec, 2)

    # abstract
    ab = doc.add_paragraph()
    run = ab.add_run("Abstract: ")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    r2 = ab.add_run(cite_sub(content.ABSTRACT))
    r2.italic = True
    r2.font.size = Pt(9)
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    kw = doc.add_paragraph()
    run = kw.add_run("Index Terms: ")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    kw.add_run(", ".join(content.KEYWORDS)).font.size = Pt(9)

    # sections, numbered with roman numerals. the back-matter sections
    # (acknowledgment, availability) stay unnumbered, as IEEE does them.
    romans = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    num = 0
    for title, paras in content.SECTIONS:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title in UNNUMBERED:
            head = title.upper()
        else:
            head = f"{romans[num]}.  {title.upper()}"
            num += 1
        run = h.add_run(head)
        run.bold = True
        run.font.size = Pt(10)
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(3)
        if title == "Results":
            add_ranking_table(doc)
            add_figure(doc)
        for p in paras:
            add_body(doc, p)

    # references
    h = doc.add_paragraph()
    run = h.add_run("REFERENCES")
    run.bold = True
    run.font.size = Pt(10)
    for i, (_key, ref) in enumerate(content.REFERENCES):
        p = doc.add_paragraph()
        p.add_run(f"[{i + 1}] {ref}").font.size = Pt(8)
        p.paragraph_format.space_after = Pt(0)

    return doc


if __name__ == "__main__":
    render().save(HERE / "icdm.docx")
    print("wrote icdm.docx")
